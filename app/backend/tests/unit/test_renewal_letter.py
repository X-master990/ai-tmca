"""續約函服務 unit tests — ROC 格式化、預填邏輯、模板渲染。

render 測試需要 repo 內的 jinja 模板；找不到就 skip（不綁特定部署環境）。
"""
from datetime import date
from pathlib import Path
from types import SimpleNamespace

import pytest

import app.services.renewal_letter as rl
from app.services.renewal_letter import (
    UNIT_FEE,
    build_prefill,
    roc_long,
    roc_short,
    _last_day_of_next_month,
    _plus_one_year_minus_day,
)

REPO_TEMPLATE = (
    Path(__file__).resolve().parents[4] / "templates" / "承辦" / "續約函模板.docx"
)


def rec(**kw):
    """輕量 Record stub（build_prefill 只讀屬性，不需 DB）。"""
    base = dict(
        id=1,
        holder_name="金紅視聽歌唱坊",
        invoice_title=None,
        action_type="新辦",
        period_start=None,
        period_end=None,
        use_address="台中市中區自由路二段135-1號10樓",
        qty=None,
    )
    base.update(kw)
    return SimpleNamespace(**base)


class TestRocFormat:
    def test_roc_long_pads_month_day(self):
        assert roc_long(date(2026, 6, 1)) == "115 年 06 月 01 日"

    def test_roc_short_no_pad(self):
        assert roc_short(date(2026, 5, 29)) == "115年5月29日"


class TestDateHelpers:
    def test_plus_one_year_minus_day(self):
        assert _plus_one_year_minus_day(date(2026, 6, 1)) == date(2027, 5, 31)

    def test_plus_one_year_leap(self):
        # 2024/3/1 起 → 2025/2/28（隔年無 2/29）
        assert _plus_one_year_minus_day(date(2024, 3, 1)) == date(2025, 2, 28)

    def test_last_day_of_next_month(self):
        assert _last_day_of_next_month(date(2026, 5, 12)) == date(2026, 6, 30)

    def test_last_day_of_next_month_december(self):
        assert _last_day_of_next_month(date(2026, 12, 15)) == date(2027, 1, 31)


class TestBuildPrefill:
    TODAY = date(2026, 5, 29)

    def test_expiring_row_computes_next_period(self):
        # 即將到期行：只有 period_end → 自次日起算一年
        pf = build_prefill(rec(period_end=date(2026, 5, 31), qty=2), today=self.TODAY)
        assert pf["period_start"] == "115 年 06 月 01 日"
        assert pf["period_end"] == "116 年 05 月 31 日"
        assert pf["qty"] == 2
        assert pf["amount"] == 2 * UNIT_FEE
        assert pf["issue_date"] == "115年5月29日"
        assert pf["pay_deadline"] == "115年6月30日"
        assert pf["recipient"] == "金紅視聽歌唱坊"

    def test_renewal_row_uses_existing_period(self):
        # 續約行：已有起迄 → 直接用，不再 +1 年
        pf = build_prefill(
            rec(
                action_type="續約",
                period_start=date(2026, 6, 1),
                period_end=date(2027, 5, 31),
                qty=None,
                use_address=None,
            ),
            today=self.TODAY,
        )
        assert pf["period_start"] == "115 年 06 月 01 日"
        assert pf["period_end"] == "116 年 05 月 31 日"
        assert pf["qty"] == 1  # 無台數 → 1
        assert pf["amount"] == UNIT_FEE
        assert pf["business_address"] == ""

    def test_renewal_row_with_only_period_end_back_computes_start(self):
        # 迴歸：續約行只有迄日（缺起日）不可被當「舊到期日」往後再滾一年
        pf = build_prefill(
            rec(
                action_type="續約",
                period_start=None,
                period_end=date(2027, 5, 31),
                qty=1,
            ),
            today=self.TODAY,
        )
        # 應回推起日 116/05/31 - 1年 + 1天 = 115/06/01，而非把迄日當到期日滾成 117 年
        assert pf["period_start"] == "115 年 06 月 01 日"
        assert pf["period_end"] == "116 年 05 月 31 日"

    def test_recipient_falls_back_to_invoice_title(self):
        pf = build_prefill(
            rec(holder_name=None, invoice_title="某某企業社"), today=self.TODAY
        )
        assert pf["recipient"] == "某某企業社"

    def test_no_period_yields_blank(self):
        pf = build_prefill(rec(period_end=None, period_start=None), today=self.TODAY)
        assert pf["period_start"] == "" and pf["period_end"] == ""


@pytest.mark.skipif(not REPO_TEMPLATE.exists(), reason="repo 模板不存在")
class TestRender:
    def test_render_substitutes_all_fields(self, monkeypatch):
        from docx import Document
        from io import BytesIO

        monkeypatch.setattr(rl, "TEMPLATE_PATH", REPO_TEMPLATE)
        data = build_prefill(
            rec(period_end=date(2026, 5, 31), qty=2), today=date(2026, 5, 29)
        )
        out = rl.render_letter(data)
        assert out and len(out) > 5000

        d = Document(BytesIO(out))
        texts = [p.text for p in d.paragraphs]
        joined = "\n".join(texts)
        assert "受文者：金紅視聽歌唱坊" in joined
        assert "發文日期：115年5月29日" in joined
        assert "★繳費期限115年6月30日" in joined
        assert "中華民國 115 年 06 月 01 日起至 116 年 05 月 31 日止。" in joined
        assert "台中市中區自由路二段135-1號10樓" in joined
        assert "申報台數： 2 台" in joined
        assert "應付金額： 7350 元" in joined

    def test_render_blank_context_no_crash(self, monkeypatch):
        monkeypatch.setattr(rl, "TEMPLATE_PATH", REPO_TEMPLATE)
        out = rl.render_letter({})  # 缺項補空字串，不應 raise
        assert out and len(out) > 5000

    def test_xml_special_chars_do_not_truncate_doc(self, monkeypatch):
        # 迴歸：店名/地址含 & < > 等字元時，autoescape 應跳脫，整份文件不可被截斷
        from docx import Document
        from io import BytesIO

        monkeypatch.setattr(rl, "TEMPLATE_PATH", REPO_TEMPLATE)
        out = rl.render_letter(
            {
                "recipient": "A & B <旗艦店>",
                "business_address": "台中市<中區> & 自由路",
                "issue_date": "115年5月29日",
                "qty": 1,
                "amount": 3675,
            }
        )
        d = Document(BytesIO(out))
        joined = "\n".join(p.text for p in d.paragraphs)
        assert "受文者：A & B <旗艦店>" in joined
        assert "台中市<中區> & 自由路" in joined
        # 文件未被截斷：模板末段的銀行帳號仍在
        assert "0851-10-0050899" in joined

    def test_user_input_is_not_re_evaluated_as_jinja(self, monkeypatch):
        # 安全：使用者輸入的 {{ }} 不應被當模板二次執行
        from docx import Document
        from io import BytesIO

        monkeypatch.setattr(rl, "TEMPLATE_PATH", REPO_TEMPLATE)
        out = rl.render_letter({"recipient": "{{ 7*7 }}", "amount": 3675})
        joined = "\n".join(p.text for p in Document(BytesIO(out)).paragraphs)
        assert "受文者：{{ 7*7 }}" in joined
        assert "受文者：49" not in joined
