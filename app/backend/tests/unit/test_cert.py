"""證書服務 unit tests — 類別對應、欄位 resolver、模板渲染。

render 測試需要 repo 內的正面模板；找不到就 skip（不綁特定部署環境）。
"""
from datetime import date
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import pytest

import app.services.cert as cert
from app.services.cert import (
    CATEGORY_TO_TEMPLATE,
    _resolve_field,
    build_cert_prefill,
    has_cert,
    render_cert,
)

REPO_CERT_DIR = Path(__file__).resolve().parents[4] / "templates" / "證書"


def rec(**kw):
    base = dict(
        category_code="COMPUTER_KARAOKE",
        cert_no="A2-115-0001",
        holder_name="金紅視聽歌唱坊",
        use_address="台中市中區自由路二段135-1號10樓",
        period_start=date(2026, 6, 1),
        period_end=date(2027, 5, 31),
        qty=2,
        brand="金嗓",
        serial_no="SN-001",
        extra={},
    )
    base.update(kw)
    return SimpleNamespace(**base)


class TestMapping:
    def test_funeral_has_no_cert(self):
        assert has_cert("FUNERAL") is False

    def test_known_categories_have_cert(self):
        for c in ("COMPUTER_KARAOKE", "SELF_SERVICE_KTV", "PUBLIC_TRANSMIT", "HALL_ROOM"):
            assert has_cert(c) is True

    def test_borrowed_templates_point_to_karaoke(self):
        # 社區管委會 / 公益伴唱機 借用電腦伴唱機證書
        assert CATEGORY_TO_TEMPLATE["PUBLIC_KARAOKE"] == CATEGORY_TO_TEMPLATE["COMPUTER_KARAOKE"]
        assert CATEGORY_TO_TEMPLATE["COMMUNITY_BOARD"] == CATEGORY_TO_TEMPLATE["COMPUTER_KARAOKE"]


class TestResolveField:
    def test_cert_no_aliases(self):
        r = rec(cert_no="X-1")
        assert _resolve_field("證號", r) == "X-1"
        assert _resolve_field("證書編號", r) == "X-1"

    def test_holder_aliases(self):
        r = rec(holder_name="某店")
        assert _resolve_field("持證人", r) == "某店"
        assert _resolve_field("持證者", r) == "某店"

    def test_address_aliases(self):
        r = rec(use_address="台北市A路1號")
        assert _resolve_field("使用地址", r) == "台北市A路1號"
        assert _resolve_field("營業地址", r) == "台北市A路1號"

    def test_roc_date_parts(self):
        r = rec(period_start=date(2026, 6, 1), period_end=date(2027, 5, 31))
        assert _resolve_field("起年", r) == "115"
        assert _resolve_field("起月", r) == "6"
        assert _resolve_field("起日", r) == "1"
        assert _resolve_field("終年", r) == "116"
        assert _resolve_field("終月", r) == "5"
        assert _resolve_field("終日", r) == "31"

    def test_extra_backed_fields(self):
        r = rec(extra={"floor_area": "120", "street_cert_no": "街-7", "platform_name": "YT", "song_count": "50"})
        assert _resolve_field("坪數", r) == "120"
        assert _resolve_field("藝人證號", r) == "街-7"
        assert _resolve_field("平台名稱", r) == "YT"
        assert _resolve_field("總曲數", r) == "50"
        assert _resolve_field("首", r) == "50"

    def test_unknown_field_blank(self):
        assert _resolve_field("每場活動場", rec()) == ""

    def test_none_values_blank(self):
        r = rec(cert_no=None, period_start=None, qty=None, brand=None, extra={})
        assert _resolve_field("證號", r) == ""
        assert _resolve_field("起年", r) == ""
        assert _resolve_field("台數", r) == ""
        assert _resolve_field("廠牌名稱", r) == ""
        assert _resolve_field("坪數", r) == ""


@pytest.mark.skipif(not REPO_CERT_DIR.exists(), reason="repo 證書模板目錄不存在")
class TestRender:
    def _dir(self, monkeypatch):
        monkeypatch.setattr(cert, "TEMPLATE_DIR", REPO_CERT_DIR)

    def test_prefill_returns_template_fields(self, monkeypatch):
        self._dir(monkeypatch)
        pre = build_cert_prefill(rec(category_code="COMPUTER_KARAOKE"))
        assert pre["template"] == "A2電腦伴唱機證書-正面.docx"
        names = {f["name"] for f in pre["fields"]}
        assert {"持證者", "營業地址", "台數", "證書編號", "起年", "終日"} <= names
        byname = {f["name"]: f["value"] for f in pre["fields"]}
        assert byname["持證者"] == "金紅視聽歌唱坊"
        assert byname["台數"] == "2"

    def test_render_fills_and_clears_placeholders(self, monkeypatch):
        self._dir(monkeypatch)
        pre = build_cert_prefill(rec(category_code="COMPUTER_KARAOKE", serial_no=""))
        vals = {f["name"]: f["value"] for f in pre["fields"]}
        out = render_cert("COMPUTER_KARAOKE", vals)
        from docx import Document

        joined = "".join(p.text for p in Document(BytesIO(out)).paragraphs)
        assert "金紅視聽歌唱坊" in joined
        assert "115年6月1日" in joined and "116年5月31日" in joined
        assert "«" not in joined  # 無殘留合併佔位符（含留空的機號）

    def test_render_strips_control_chars(self, monkeypatch):
        # 迴歸：含 NULL/控制字元的值不可讓 lxml 丟 ValueError；清掉後正常產出
        self._dir(monkeypatch)
        out = render_cert(
            "COMPUTER_KARAOKE",
            {
                "證書編號": "A2-1",
                "持證者": "金\x00紅\x07視聽",  # 內含 NULL 與控制字元
                "營業地址": "台北市\x1f測試路",
                "起年": "115", "起月": "6", "起日": "1",
                "終年": "116", "終月": "5", "終日": "31",
            },
        )
        from docx import Document

        joined = "".join(p.text for p in Document(BytesIO(out)).paragraphs)
        assert "金紅視聽" in joined  # 控制字元被清掉、字面保留
        assert "\x00" not in joined and "\x07" not in joined

    def test_render_unknown_category_raises(self, monkeypatch):
        self._dir(monkeypatch)
        with pytest.raises(ValueError):
            render_cert("FUNERAL", {})

    def test_render_missing_template_raises(self, monkeypatch):
        monkeypatch.setattr(cert, "TEMPLATE_DIR", Path("/nonexistent/dir"))
        with pytest.raises(FileNotFoundError):
            render_cert("COMPUTER_KARAOKE", {})
